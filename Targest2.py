# This program will search for red colored tags to get the information it needs
#
# Instructions on how to use the program:
# 1. run program
# 2. click on button "Choose document" and choose a document
# 3. Do step 2 as many times as you want, to uadd as many documents as you like
# 4. After you are done choosing documents, click on the GenerateReport button
# 5. Then you can click on the "open generated report" button, which will automatically
# open up your word document report created from your documents
# 6. When you are done, click "End Program"

# from debug import debug
import logging
# import pdb
import docx
from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
import tkinter as tk
from tkinter import *
from tkinter import filedialog
from typing import Tuple

from tkinter import scrolledtext
import re
import copy
import time

# This libraries are for opening word document automatically
import os
import platform
import subprocess

# This library is for opening excel document automatically
import xlwings as xw
import pandas as pd
import matplotlib.pyplot as plt

import Gui

from tkinter import messagebox


# Set up the logger for catching errors
logging.basicConfig(level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')

logger = logging.getLogger(__name__)

global report1
report1 = Document()
report1.add_heading('All Tags in each document', 0) #create word document
global paragraph0 
paragraph0 = report1.add_paragraph()
report1.save('ChildandParentTagsTables.docx')


global report3
report3 = Document()
report3.add_heading('All Tags and Requirement Tracing', 0) #create word document
global paragraph 
paragraph = report3.add_paragraph()
report3.save('AllChildandParentTags.docx')

global orphanReport
orphanReport = Document()
orphanReport.add_heading('Orphan Report', 0)
global paragraph2
paragraph2 = orphanReport.add_paragraph()
runnerOrphan = paragraph2.add_run("These are the orphan tags that were found in the documents: ")
runnerOrphan.bold = True  # makes it bold
orphanReport.save('orphanReport.docx')

global childlessReport
childlessReport = Document()
childlessReport.add_heading('Childless Report', 0)

global paragraph3
paragraph3 = childlessReport.add_paragraph()
runnerOrphan = paragraph3.add_run("These are the childless tags that were found in the documents: ")
runnerOrphan.bold = True  # makes it bold
childlessReport.save('ChildlessTags.docx')

global TBVReport
TBVReport = Document()
TBVReport.add_heading('TBV Tags', 0) #create word document
global paragraph5 
paragraph5 = TBVReport.add_paragraph()
TBVReport.save('TBVReport.docx')

global TBDReport
TBDReport = Document()
TBDReport.add_heading('TBD Tags', 0) #create word document
global paragraph6 
paragraph6 = TBDReport.add_paragraph()
TBDReport.save('TBDReport.docx')

global dicts2Copy # This will hold the dicts2 content in all documents
dicts2Copy = {}

global parents2Copy # parents2 list copy
parents2Copy = []

global filtered_L # Will store the ones without a child tag
filtered_L = []

global filtered_LCopy
filtered_LCopy = []

global fullText2Copy
fullText2Copy = []

global parents2 #list of parent tags or child tags
parents2 = []

# creates a dict for parent and child tags
global dicts
dicts = {}

global OrphanChild2
OrphanChild2 = []
global OrphanChild2Copy
orphanChildren2Copy = []

global dicts10
dicts10 = {}
global dicts3
dicts3 = {}  # will hold parentTag and text, Orphan tags
global dicts2
dicts2 = {}  # will hold parentTag and text
global orphanDicts
orphanDicts = {}  # orphan dictionary

global parents9
parents9 = []

# declaring different lists that will be used to store, tags and sentences
global parentTags
parentTags = []
global parent  # This will be used to store everything
parent = []
global child # Used to Store child tags
child = [] 
global noChild # Used to Store parentTags with no child
noChild = []  
global withChild # Used to Store parentTags with child tag
withChild = [] 
global parents # Will be used for future function
parents = [] 

global orphanTagText
orphanTagText = []  # Will be used to hold text of orphanChildTags


# reads the text in the document and use the getcoloredTXT function to get the colored text
def readtxt(filename, color: Tuple[int, int, int]):
    try:
        doc = docx.Document(filename)
        text10 = ""
        fullText = []
        new = []
        global everything
        everything = []  # list of tags and text

        for para in doc.paragraphs:
            # Getting the colored words from the doc
            if (getcoloredTxt(para.runs, color)):
                # Concatenating list of runs between the colored text to single a string
                sentence = "".join(r.text for r in para.runs)
                if len(sentence) > 5:  # this chcekcts if sentence has atleast 5 characters
                    fullText.append(sentence)
                #print(sentence) # Prints everything in the terminal
                everything.append(sentence)
                text10 = sentence
                parent.append("".join(r.text for r in para.runs))

        #print(fullText)
        global hasChild # Will store the ones with a child tag
        global fullText2 # will store everything found
        global children

        global orphanss
        orphanss = []

        global orphanChildren2 # Will store the orphan child tags for orphanReport
        orphanChildren2 = []
        # Finds the lines without a parentTag
        filtered_L = [value for value in fullText if "[" not in value]

        filtered_L = [s.replace(": ", ":") for s in filtered_L]
        # Finds the lines with a parentTag
        filtered_LCopy.extend(filtered_L)
        hasChild = [value for value in fullText if "[" in value]
        # will store everything found
        fullText2 = [value for value in fullText]
        fullText2 = [s.replace(": ", ":") for s in fullText2]
        fullText2 = [s.replace("[ ", "[") for s in fullText2]
        fullText2 = [s.replace("] ", "]") for s in fullText2]
        fullText2Copy.extend(fullText2)
        
        return fullText, filtered_L, hasChild, filtered_LCopy, fullText2Copy, fullText2

    except Exception as e:
    # Log an error message
        logging.error('readtxt(): ERROR', exc_info=True)


def getcoloredTxt(runs, color): 
    coloredWords, word = [], ""
    try:
        for run in runs:
            if run.font.color.rgb == RGBColor(*color):
                word += str(run.text) # Saves everything found

            elif word != "":  # This will find the parentTags
                coloredWords.append(word)
                parentTags.append(word)
                parents.append(word)
                word = ""

        if word != "":  # This will find the parentTags
            coloredWords.append(word + "\n")
            # word = removeAfter(word)
            child.append(word)
            withChild.append(word)

            

    except Exception as e:
        logging.error('getColoredText(): ERROR', e)
    else:
        # Log a success message
        logging.info('getColoredText(): PASS')

    return coloredWords # returns everything found


def generateReport(): #Will generate the report for tags
    try:
        global filepath
        global filepath2
        # create a similar method for opening a folder
        filepath = filedialog.askopenfilename(initialdir="/",
                                            title="",
                                            filetypes = (("text file", "*.txt"),
                                                        ("all files","*.*")))
        file = open(filepath,'r')
        #print(filepath)
        file.close()
        # Will store the filepath to the document as a string
        filepath2 = str(filepath)
        a = (filepath2)
        with open(a) as file_in:
            lines = []
            for line in file_in:
                lines.append(line)
        for line2 in lines:
            print(line2)
            line3 = str(line2)
            line4 = line3.replace('\\', '/')
            line5 = line4.replace('"', '')
            line6 = line5.replace("\n", "")
            print(line6)
            fullText = readtxt(filename=line6,
                            color=(255, 0, 0))
            print(line4)

            #filtered_L = readtxt(filename=filepath2, #For future use
            #                   color=(255, 0, 0))
            fullText10 = str(fullText)
            s = ''.join(fullText10)
            w = (s.replace (']', ']\n\n'))
            #paragraph = report3.add_paragraph()
            paragraph0 = report1.add_paragraph()
            #paragraph2 = orphanReport.add_paragraph()
            filepath3 = str(line4.rsplit('/', 1)[-1]) # change filepath to something.docx
            filepath3 = filepath3.split('.', 1)[0] # removes .docx of the file name
            print(filepath3 + " added to the report")
            nameOfDoc = (filepath3 + " added to the report\n")
            Gui.Txt.insert(tk.END, nameOfDoc) #print in GUI (m = main.py)
            #runner = paragraph.add_run("\n" + "Document Name: " + filepath3 + "\n")
            #runner.bold = True  # makes the header bold

            runner0 = paragraph0.add_run("\n" + "Document Name: " + filepath3 + "\n")
            runner0.bold = True  # makes the header bold

            # w will be used in the future
            w = (w.replace ('([', ''))
            w = (w.replace (',', ''))
            w = (w.replace ('' '', ''))

            # creates a table for report 1
            #table = report3.add_table(rows=1, cols=2)

            # creates a table for report 3
            table1 = report1.add_table(rows=1, cols=2)


            row1 = table1.rows[0].cells
            row1[0].text = 'Front Tag'
            row1[1].text = 'Back Tag/tags'

            # Adding style to a table
            #table.style = 'Colorful List'

            table1.style = 'Colorful List'

            # Now save the document to a location
            #report3.save('report.docx')

            report1.save('ChildandParentTagsTables.docx')

            #orphanReport.save('orphanReport.docx')
            # Adds headers in the 1st row of the table


            e = 0

            child2 = removeAfter(child) #removes everything after the parent tag if there is anything to remove
            # while loop until all the  parentTags has been added to the report


            parents2 = copy.deepcopy(parentTags) # copy of parent tags list
            parents2Copy.extend(parents2)
            childCopy = copy.deepcopy(child2)
            noParent = []
            noParent2 = []
            global orphanChild
            orphanChild = []
            orphanChildParent = []
            parents9000 = []

            parents2 = [s.replace(" ", "") for s in parents2] # gets rid of space
            while parentTags:
                #row = table.add_row().cells # Adding a row and then adding data in it.
                #row[0].text = parentTags[0] # Adds the parentTag to the table
                #report3.add_paragraph(parentTags[0])
                row1 = table1.add_row().cells # Adding a row and then adding data in it.
                

                row1[0].text = parentTags[0] # Adds the parentTag to the table
                
                
                noParent.append(parentTags[0])


                #for child100 in child2:
                 #   report3.add_paragraph(child100)
                for ch in child2:
                    if "[" not in ch:
                        child2.remove(ch)
                    
                           

                if e < len(fullText2):  #as long as variable e is not higher than the lines in fullText2
                    if fullText2[e] in filtered_LCopy: #filtered_L contains the child tags without a parent tag
                        #report3.add_paragraph(parentTags[0] + " has no child tag")
                        orphanChild.append(parentTags[0])
                        orphanChildren2.append(parentTags[0])
                        #orphanReport.add_paragraph(parentTags[0] + " has no child tag")
                        parentTags.remove(parentTags[0]) # Removes that tag after use
                        noParent2.append(" ")
                        parents9000.append(" ")
                        orphanChildParent.append(" ")
                        #row[1].text = " " # No parent tag, so adds empty string to that cell
                        if child2:
                            if "[" not in child2[0]:
                                row1[1].text = " " # No parent tag, so adds empty string to that cell
                        if child2:
                            if "[" not in child2[0]: # if it is not a parent tag
                                child2.remove(child2[0])  # Removed that tag from the list

                        
                        #if child2:
                         #   child2.remove(child2[0])  # Removed that tag from the list

                        e += 1

                    elif fullText2[e] not in filtered_LCopy:
                        parentTags.remove(parentTags[0]) # Removes that tag after use
                        if child2:
                            #row[1].text = child2[0] #Adds childTag to table
                            
                            if "[" not in child2[0]:
                                child2.remove(child2[0])  # Removed that tag from the list
                                #report3.add_paragraph(child2[0])
                            #report3.add_paragraph(child2[0])
                            row1[1].text = child2[0] #Adds childTag to table                        
                            parents9000.append(child2[0])
                            noParent.append(child2[0])
                            child2.remove(child2[0])  # Removed that tag from the list
                            e += 1

            parents9.extend(parents9000)
            orphanChildren2Copy.extend(orphanChildren2)

            # Make sure everything is cleared before the program gets the next document
            child2.clear()
            parentTags.clear()
            child.clear()
            #report3.save('report.docx') #Saves in document "report3"
            orphanReport.save('orphanReport.docx') #Saves in document "orphanReport"
            report1.save('ChildandParentTagsTables.docx') #Saves in document "report3"

            global dicts11
            dicts11 = dict(zip(parents2, childCopy)) #creates a dictionary if there is a child tag and parent tag
            dicts.update(dicts)

            noParent = [s.replace(" ", "") for s in noParent]
            #dicts3 = dict(zip(noParent, noParent2)) # dictionary for parent tags without child tags
            orphanChild = [s.replace(" ", "") for s in orphanChild]

            
            #orphanChildren2Copy = copy.deepcopy(orphanChildren2) # copy of orphanChildren2 list

            dicts9000 = dict(zip(orphanChild, orphanChildParent)) # orphan dictionary
            orphanDicts.update(dicts9000)
            OrphanChild2.extend(orphanChild)

            text2 = removeParent(everything) # child tag and text
            # print(text2)
            text3 = removechild(text2)  # only text list
            # print(text3)
            text4 = removeText(text2) # child tags
            # print(text4) #only parent tag list
            #text8 = [s.replace(" ", "") for s in text4]

            parents9000 = [x.strip(' ') for x in parents9000]
            #dicts3 = dict(zip(parents2, childCopy))
            dicts3 = dict(zip(parents2, parents9000))
            dicts10.update(dicts3)
            dicts2 = dict(zip(parents2, text3)) # creates a dictionary with child tags and text
            dicts100 = copy.deepcopy(dicts2)
            
            sorted(dicts2.keys()) # sorts the keys in the dictionary
            dicts2Copy.update(dicts100)



            toggle_state2() # This will enable the generate report button
            toggle_state6() # This will enable the open allTags report button
            # for tg in parents2:
            #     if "TBV:" in tg:
            #         TBVReport.add_paragraph(tg)
            # TBVReport.save('TBVReport.docx')
            
        return filepath2, filtered_L, orphanChild
        return parents2, dicts2, dicts10, dicts2Copy, parents2Copy, fullText2, filtered_LCopy, dicts3, orphanDicts, OrphanChild2
        
    except Exception as e:
        # Log an error message
        logging.error('generateReport(): ERROR', e)
    else:
        # Log a success message
        logging.info('generateReport(): PASS')


def generateReport2():
    try:

        #print("here is dicts10 before:")
        #print(dicts10)
        global dicts11111 # Will be used for the excel report later for child - parent
        dicts11111 = {}
        dicts11111 = copy.deepcopy(dicts10)

        # counters for Excel report2

        global counter1
        counter1 = 2
        global counter2
        counter2 = 1
        global counter3
        counter3 = 0
        global cell
        cell = 0;
        global cell2
        cell2 = 0;



        #excelReport2.range("A3").value = 'Childless Tags'

        pattern = r'\[([^\]]+)\]'  
        for key in dicts10:
                if type(dicts10[key]) == str:
                    matches2 = re.findall(pattern, (dicts10[key])) 
                if len(matches2) > 1:
                        
                        
                        #print("There is more than one ']' in the input string.", dicts10[key] )
                        dicts10[key] = []
                        parents2 = []
                        for match in matches2:
                            parents2.append(match)
                            
                        for tag in parents2:
                            #parentChild2.setdefault("[PUMP:SRS:1]", ["[PUMP:PRS:0]"]).append("[PUMP:PRS:2]")
                            #parentChild2.setdefault(key, [parentChild2[key]]).append(tag)
                            #parentChild2[key].append(tag)
                            tag = (tag.replace(' ', ''))
                            dicts10[key] += [tag]
                            #parentChild2.setdefault(key, ["[PUMP:PRS:0]"]).append(tag)
                            #parentChild2.setdefault(key, [parentChild2[key]]).append(match)
                            
                else:
                    #print("There is only one ']' in the input string.")
                    print("")
        
        #print("here is dicts10 after:")
        #print(dicts10)


        #report3.add_paragraph("all parents:") # header for all parents

        parents10 = [] # list of all the parent tag tags
        for value11 in dicts10.values():
            # if the value is a list, extend the parents list with the list
            if isinstance(value11, list):
                parents10.extend(value11)
                
            # if the value is not a list, append the value to the parents list
            else:
                parents10.append(value11)

        # for loop to add the parents to the report
        #for parent in parents10:
         #   report3.add_paragraph(parent)
        

        # create a list of all the keys in the dictionary (all child tags)
        values_list = list(dicts2Copy.keys())

        #  creates a list of all the child tags that are not in the parents list
        global childless
        childless = [] 


        global TBVTags
        TBVTags = []

        global TBDTags
        TBDTags = []

        # for loop to check if the child tag is in the parents list
        for element in values_list:
            if "".join(element) not in "".join(parents10):
                childless.append(element)


        # sorts the childless list
        childless.sort()
        
        # for loop to add the childless tags to the report
        for child0 in childless:
            childlessReport.add_paragraph(child0) 

        childlessReport.save('ChildlessTags.docx') #Saves in document "ChildlessTags.docx"

        

        # declaring counters
        m = 0
        k = 0
        i = 0
        o = 0
        z = 0

        orphanTagText = removechild(filtered_LCopy)

        #dict(sorted(dicts2Copy.items(), key=lambda item: item[1])) # sorts by value/parentTag, not working at the moment
        #print(parents2Copy)
        #print(dicts10)
        #print(dicts2Copy)
        #print(filtered_LCopy)
        #print(fullText2Copy)
        #print(parents2Copy)
        #print(orphanTagText)
        #print(dicts2Copy)
        #report3.add_paragraph("\n") # Adds a line space from the table
        while m < len(dicts2Copy):
            #print(m)
            #if fullText2Copy[k] not in filtered_LCopy:
            if z < len(dicts2Copy) and dicts2Copy:
                z += 1
                duplicates = []
                for key, value in dicts2Copy.items():
                    
                    m += 1
                    if k < len(fullText2Copy) and fullText2Copy[k] not in filtered_LCopy:
                        #for key, value in dicts2Copy.items() and key, value in dicts3.items(): #work on this here and try
                        #report3.add_paragraph("\n")
                        stringKey = str(key)
                        stringKey2 = (stringKey.replace(' ', ''))
                        if str(stringKey2) in dicts10: # if the key is in the dictionary
                            text = dicts10[str(stringKey2)]
                        
                        

                        if isinstance(text, list):
                            #print("it is a list")
                            #report3.add_paragraph("List tags found") # display the parent tag, included brackets
                            
                            for tag in text:
                                #tag = ("[" + tag)
                                PTags = tag.split(']')
                                PTags = [s.strip() + ']' for s in PTags]
                                tag.strip()
                            
                                #report3.add_paragraph(tag)
                                if (str(tag) in duplicates):
                                    #print("in duplicates")
                                    print("")
                                    
                                    
                                    

                                else:
                                    parentTag1 = ('['+tag+']')
                                    
                                    #cell = str('A'+ str(counter1))
                                    #cell2 = str(str(parentTag1))
                                    #xcelReport2.range(cell).value = cell2
                                    counter1 += 2 # counter for excel report
                                    counter2 += 1 # counter for excel report

                                    #wb2.save('AllTags.xlsx') # Saving excel report as 'AllTags.xlsx'
                                    if "TBV:" in parentTag1:
                                        TBVReport.add_paragraph(parentTag1)
                                        TBVReport.save('TBVReport.docx')
                                        TBVTags.append(parentTag1)
                                    if "TBD:" in parentTag1:
                                        TBDReport.add_paragraph(parentTag1)
                                        TBDReport.save('TBDReport.docx')
                                        TBDTags.append(parentTag1)
                                        
                                    report3.add_paragraph(parentTag1)
                                    tag.strip()
                                    duplicates.append(str(tag))
                                    
                                    

                                    for x in PTags:
                                        #report3.add_paragraph(x)

                                        keyCheck = (x.replace('[', ''))
                                        keyCheck2 = (keyCheck.replace(']', ''))
                                        keyCheck3 = (keyCheck2.replace(']', ''))
                                        keyCheck4 = (keyCheck3.replace(' ', ''))
                                        keyCheck4.split()
                                                                
                                        

                                        if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                            if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                                if "TBV:" in parentTag1:
                                                    TBVReport.add_paragraph(dicts2Copy[str(keyCheck4)])
                                                    TBVReport.save('TBVReport.docx')
                                                if "TBD:" in parentTag1:
                                                    TBDReport.add_paragraph(dicts2Copy[str(keyCheck4)])
                                                    TBDReport.save('TBDReport.docx')

                                            
                                                report3.add_paragraph(dicts2Copy[str(keyCheck4)])

                                                
                                            
                                            #orphanReport.add_paragraph(dicts2Copy[str(keyCheck4)])

                                        else:
                                            report3.add_paragraph("Requirement text not found")
                                            orphanChildren2Copy.append(str(keyCheck4))
                                            #orphanReport.add_paragraph("Requirement text not found")
                                        #print(dicts10[str(key)])
                                        #report3.add_paragraph(dicts10[str(stringKey)])
                                        #for dicts10[str(stringKey)] in dicts10:
                                        #report3.add_paragraph(dicts10[str(stringKey)])
                                        for b in PTags:
                                            b = (b.replace(']', ''))
                                            #report3.add_paragraph("I'm b")
                                            #report3.add_paragraph(b)
                                            #report3.add_paragraph("I'm tag")
                                            #report3.add_paragraph(tag)


                                            if b == tag:
                                                i += 1
                                                hx = tag
                                                #report3.add_paragraph("I'm here")
    
                                                keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags
                                                #report3.add_paragraph("I'm keys")
                                                #report3.add_paragraph(keys)
                                                k += 1
                                                for item in keys: #keys are child tags of hx/the parent tag

                                                    if item != "" and item!= " ":
                                                        report3.add_paragraph(item, style='List Bullet')
                                                        para = report3.add_paragraph(dicts2Copy[str(item)])
                                                        para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                        if "TBD:" in parentTag1:
                                                            TBDReport.add_paragraph(item, style='List Bullet')
                                                            para3 = report3.add_paragraph(dicts2Copy[str(item)])
                                                            para3.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                            TBDReport.save('TBDReport.docx')
                                                        
                                                        
                                                    
                                                        if "TBV:" in parentTag1:
                                                            TBVReport.add_paragraph(item, style='List Bullet')
                                                            if str(item) in dicts2Copy:
                                                                para2 = TBVReport.add_paragraph(dicts2Copy[str(item)])
                                                                para2.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                            #TBVReport.add_paragraph("here1")
                                                            TBVReport.save('TBVReport.docx')
                                                            stringKey = str(item)
                                                            stringKey2 = (stringKey.replace(' ', ''))
                                                            duplicates2 = []

                                                        if "TBD:" in item:
                                                            TBDReport.add_paragraph(item)
                                                            TBDReport.save('TBDReport.docx')

                                                        if "TBV:" in item:
                                                            TBVReport.add_paragraph(item)
                                                            TBVReport.save('TBVReport.docx')

                                                            if str(stringKey2) in dicts10: # if the key is in the dictionary
                                                                #for key in dicts10:
                                                                    if item in dicts10:
                                                                        hx = item
                                                                                        #report3.add_paragraph("I'm here")
        
                                                                        keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags
                                                                                        #report3.add_paragraph("I'm keys")
                                                                                        #report3.add_paragraph(keys)
                                                                        k += 1
                                                                        for item1 in keys: #keys are child tags of hx/the parent tag

                                                                            if item1 != "" and item1!= " ":
                                                                                if item1 not in duplicates2:
                                                                                    duplicates2.append(item1)
                                                                                    tag2 = TBVReport.add_paragraph(item1)
                                                                                    tag2.paragraph_format.left_indent = Inches(0.50) # adds indentation of text
                                                                                    para = TBVReport.add_paragraph(dicts2Copy[str(item1)])
                                                                                    para.paragraph_format.left_indent = Inches(0.50) # adds indentation of text
                                                                                    TBVReport.save('TBVReport.docx')
                                                                        
                                                                
                                                                    TBVReport.save('TBVReport.docx')
                                                        
                                                        
                                                        
                                                        counter2 = counter1 - 1
                                                        cell = str('B'+ str(counter2))
                                                        cell2 = str(item)
                                                        #excelReport2.range(cell).value = cell2
                                                        counter2 += 1
                                                        counter1 += 1


                                                        #wb2.save('AllTags.xlsx') # Saving excel report as 'AllTags.xlsx'

                                                #wb2.save('AllTags.xlsx') # Saving excel report as 'AllTags.xlsx'
                                                report3.add_paragraph("\n")
                                                counter2 += 1
                                                counter1 += 1
                                                #excelReport2.autofit()



                        else:
                            #print("not a list")    
                            PTags = text.split(']')
                            PTags = [s.strip() + ']' for s in PTags]
                            PTags.pop()
                            hx10 = text
                            hx10 = hx10.replace('[', '')
                            hx10 = hx10.replace(']', '')
                            #tag.strip()
                            if (str(hx10) in duplicates):
                                #print("in duplicates")
                                print("")
                                    

                            else:
                                #parentTag1 = ('['+tag+']')
                                #report3.add_paragraph(parentTag1)
                                #tag.strip()
                                

                                for x in PTags:
                                    
                                    #report3.add_paragraph(str(text))
                                    keyCheck = (x.replace('[', ''))
                                    keyCheck2 = (keyCheck.replace(']', ''))
                                    keyCheck3 = (keyCheck2.replace(']', ''))
                                    keyCheck4 = (keyCheck3.replace(' ', ''))
                                    report3.add_paragraph(x) # display the parent tag, included brackets
                                    if "TBV:" in x:
                                                            TBVReport.add_paragraph(x)
                                                            TBVReport.save('TBVReport.docx')
                                                            TBVTags.append(x)
                                    if "TBD:" in x:
                                                            TBDReport.add_paragraph(item)
                                                            TBDReport.save('TBDReport.docx')
                                                            TBDTags.append(x)
                                                            
                                    cell = str('A'+ str(counter1))
                                    cell2 = str(str(x))
                                    #excelReport2.range(cell).value = cell2
                                    counter1 += 2 # counter for excel report
                                    counter2 += 1 # counter for excel report

                                    #wb2.save('AllTags.xlsx') # Saving excel report as 'AllTags.xlsx'

                                    if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                        if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                            report3.add_paragraph(dicts2Copy[str(keyCheck4)])
                                            if "TBD:" in keyCheck4:
                                                            TBDReport.add_paragraph(dicts2Copy[str(keyCheck4)])
                                                            TBDReport.save('TBDReport.docx')
                                            if "TBV:" in keyCheck4:
                                                            TBVReport.add_paragraph(dicts2Copy[str(keyCheck4)])
                                                            TBVReport.save('TBVReport.docx')
                                        #orphanReport.add_paragraph(dicts2Copy[str(keyCheck4)])

                                    else:
                                        report3.add_paragraph("Requirement text not found")
                                        orphanChildren2Copy.append(str(keyCheck4))
                                        #orphanReport.add_paragraph("Requirement text not found")
                                    #print(dicts10[str(key)])
                                    #report3.add_paragraph(dicts10[str(stringKey)])
                                    #for dicts10[str(stringKey)] in dicts10:
                                    #report3.add_paragraph(dicts10[str(stringKey)])
                                    for b in PTags:


                                        if b == dicts10[str(stringKey2)]:
                                            i += 1
                                            text.strip()
                                            
                                            hx = text
                                            hx = hx.replace('[', '')
                                            hx = hx.replace(']', '')
                                            
                                            duplicates.append(str(hx))
                                            #report3.add_paragraph(str(hx))
                                            
                                            keys = [h for h, v in dicts10.items() if check_string(hx, v)]
                                            # finds all the child tags
                                            #print(keys)
                                            k += 1

                                            #if keys:
                                            #   print("list is not empty")
                                            #else:
                                            #    report3.add_paragraph("It is an orphan tag")

                                            for item in keys: #keys are child tags of hx/the parent tag

                                                if item != "" and item!= " ":
                                                    report3.add_paragraph(item, style='List Bullet')
                                                    para = report3.add_paragraph(dicts2Copy[str(item)])
                                                    para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                    if "TBD:" in x:
                                                            TBDReport.add_paragraph(item, style='List Bullet')
                                                            para = TBDReport.add_paragraph(dicts2Copy[str(item)])
                                                            para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                            TBDReport.save('TBDReport.docx')
                                                    if "TBV:" in x:
                                                            TBVReport.add_paragraph(item, style='List Bullet')
                                                            if str(item) in dicts2Copy:
                                                                para2 = TBVReport.add_paragraph(dicts2Copy[str(item)])
                                                                para2.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                            #TBVReport.add_paragraph("here1")
                                                            TBVReport.save('TBVReport.docx')
                                                            stringKey = str(item)
                                                            stringKey2 = (stringKey.replace(' ', ''))
                                                            duplicates2 = []
                                                        
                                                    if "TBD:" in item:
                                                            TBDReport.add_paragraph(item)
                                                            TBDReport.save('TBDReport.docx')

                                                    if "TBV:" in item:
                                                            TBVReport.add_paragraph(item)
                                                            TBVReport.save('TBVReport.docx')
                                                            
                                                            if str(stringKey2) in dicts10: # if the key is in the dictionary
                                                                #for key in dicts10:
                                                                    if item in dicts10:
                                                                        hx = item
                                                                                        #report3.add_paragraph("I'm here")
        
                                                                        keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags
                                                                                        #report3.add_paragraph("I'm keys")
                                                                                        #report3.add_paragraph(keys)
                                                                        k += 1
                                                                        for item1 in keys: #keys are child tags of hx/the parent tag

                                                                            if item1 != "" and item1!= " ":
                                                                                if item1 not in duplicates2:
                                                                                    duplicates2.append(item1)
                                                                                    tag2 = TBVReport.add_paragraph(item1)
                                                                                    tag2.paragraph_format.left_indent = Inches(0.50) # adds indentation of text
                                                                                    para = TBVReport.add_paragraph(dicts2Copy[str(item1)])
                                                                                    para.paragraph_format.left_indent = Inches(0.50) # adds indentation of text
                                                                                    TBVReport.save('TBVReport.docx')
                                                                        
                                                                
                                                                    TBVReport.save('TBVReport.docx')
                                                    
                                                    
                                                    counter2 = counter1 - 1
                                                    cell = str('B'+ str(counter2))
                                                    cell2 = str(item)
                                                    #excelReport2.range(cell).value = cell2
                                                    counter2 += 1
                                                    counter1 += 1

                                            report3.add_paragraph("\n")
                                            counter2 += 1
                                            counter1 += 1
                                            #excelReport2.autofit()

                            #report3.add_paragraph("\n") # Adds a line space
                            #print(k)
                            #print(m)
                            #report3.add_paragraph(key, style='List Bullet')
                            #para = report3.add_paragraph(value)
                            #para.paragraph_format.left_indent = Inches(0.25) # adds indentation ot text
                            #stringKey = dicts2Copy[str(key)]
                            #stringKey2 = (stringKey.replace(' ', ''))
                        #if k < len(fullText2Copy):
                    #elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                    elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                        k += 1
                        #report3.add_paragraph("\n")
                        if i < len(parents2Copy):
                            #report3.add_paragraph(parents2Copy[i])
                            #orphanReport.add_paragraph(parents2Copy[i])
                            #paragraph2.add("\n" +parents2Copy[i] + " Is an orphanTag" + "\n")
                            
                            #if "TBV:" in parentTag1:
                            #    TBVReport.add_paragraph(item, style='List Bullet')
                            #    TBVReport.save('TBVReport.docx')
                           # if "TBV:" in parentTag1:
                            #    TBVReport.add_paragraph(item, style='List Bullet')
                            


                            TBVReport.save('TBVReport.docx')
                            #print(parents2Copy[i])
                            #print(orphanTagText[o])
                            print("")
                            #orphanss.append(parents2Copy[i])
                        if o < len(orphanTagText):
                            #report3.add_paragraph(orphanTagText[o])
                            #orphanReport.add_paragraph(orphanTagText[o])
                            print("")
                        o += 1
                        if i < len(parents2Copy):
                            print("")
                            #orphanReport.add_paragraph(parents2Copy[i] + " is an orphan tag")
                        #m += 1

                        i += 1
                        #del dicts2Copy[list(dicts2Copy.keys())[0]] # deletes the first item in dicts2Copy

        
        # Description, this function is removing some linees, but not all, maybe need more work
        # We check if the length of the paragraph text is less than 4 using the len() function, 
        # and if so, we remove the paragraph using the _element.clear() method. Finally
        # iterate over all the paragraphs in the document
        #for i in range(len(report3.paragraphs)):
            # check if the paragraph contains less than 4 characters or a string of length less than 4
         #   if len(report3.paragraphs[i].text) < 4:
                # remove the paragraph
          #      report3.paragraphs[i]._element.clear()
                # save the modified document
           #     report3.save('AllChildandParentTags.docx')

        #if parents2Copy:
        #    for tg in parents2Copy:
        #        if "TBV:" in tg:
        #            TBVReport.add_paragraph(tg)
        #            TBVReport.save('TBVReport.docx')

        #wb2.save('AllTags.xlsx')
        msg1 = ("\nYou can now open up your Word and Excel reports\n")
        Gui.Txt.insert(tk.END, msg1) #print in GUI
        msg2 = ("\nDifferent reports include:\n")
        Gui.Txt.insert(tk.END, msg2) #print in GUI
        msg3 = ("a. All of the Tags found organized in tables\n")
        Gui.Txt.insert(tk.END, msg3) #print in GUI
        msg4 = ("b. All Child and Parent Tags\n")
        Gui.Txt.insert(tk.END, msg4) #print in GUI
        msg5 = ("c. Orphan Tags\n")
        Gui.Txt.insert(tk.END, msg5) #print in GUI
        msg6 = ("d. Childless Tags\n")
        Gui.Txt.insert(tk.END, msg6) #print in GUI
        msg7 = ("e. TBV Tags\n")
        Gui.Txt.insert(tk.END, msg7) #print in GUI
        msg8 = ("f. TBD Tags\n")
        Gui.Txt.insert(tk.END, msg8) #print in GUI
        msg9 = ("g. Tags and Requirements Excel report\n")
        Gui.Txt.insert(tk.END, msg9) #print in GUI
        msg10 = ("h. Relationship Trees Excel Report\n")
        Gui.Txt.insert(tk.END, msg10) #print in GUI
        report3.save('AllChildandParentTags.docx')
        toggle_state() #This will enable the getDoc button
        Gui.genRep.config(state="disabled") # This will disable the generate report button
        TBVReport.save('TBVReport.docx') # saves the TBV report
        toggle_state3() # this will re-enable excel report button
        #toggle_state5() # This will enable the generate orphan report button
        toggle_state7() #This will enable the getChildless document button
        toggle_state9() # This will enable the get TBV button
        toggle_state10() # This will enable the get TBD button
        toggle_state8() # This will enable the getExcel2 report button
        toggle_state13() # This will enable the guiTree View button
        orphanGenReport()
 

        
    except Exception as e:
        # Log an error message
        logging.error('generateReport2(): ERROR', e)
    else:
        # Log a success message
        logging.info('generateReport2(): PASS')

    """
        elif not dicts2Copy: # this is for orphan tags
            dict3 = dict(dicts2.items() - dicts3.items())
            for key, value in dicts3.items():
                report3.add_paragraph("\n")
                report3.add_paragraph(key)
                report3.add_paragraph(value)
                report3.add_paragraph(key + " is an orphan tags")
                m += 1
    """

def orphanGenReport():
    duplicates = []
    try:
        """
        # declaring counters
        m = 0
        k = 0
        i = 0
        o = 0
        z = 0

        orphanTagText = removechild(filtered_LCopy)
        while m < len(dicts2Copy):
            #print(m)
            #if fullText2Copy[k] not in filtered_LCopy:
            if z < len(dicts2Copy) and dicts2Copy:
                z += 1

                for key, value in dicts2Copy.items():
                    #orphanReport.add_paragraph("\n")
                    m += 1
                    if k < len(fullText2Copy) and fullText2Copy[k] not in filtered_LCopy:
                        #for key, value in dicts2Copy.items() and key, value in dicts3.items(): #work on this here and try
                        #report3.add_paragraph("\n")
                        stringKey = str(key)
                        stringKey2 = (stringKey.replace(' ', ''))
                        if stringKey2 in dicts10:
                            text = dicts10[str(stringKey2)]
                        

                        if isinstance(text, list):
                            #print("it is a list")
                            print("")
                            #report3.add_paragraph("List tags found") # display the parent tag, included brackets
                            
                            for tag in text:
                                #tag = ("[" + tag)
                                PTags = tag.split(']')
                                PTags = [s.strip() + ']' for s in PTags]
                                tag.strip()
                                if (str(tag) in duplicates):
                                    #print("in duplicates")
                                    print("")
                                    

                                else:
                                    
                                    #report3.add_paragraph(tag)
                                    #tag.strip()
                                    duplicates.append(str(tag))
                                    
                  
                                    for x in PTags:
                                        #report3.add_paragraph(x)

                                        keyCheck = (x.replace('[', ''))
                                        keyCheck2 = (keyCheck.replace(']', ''))
                                        keyCheck3 = (keyCheck2.replace(']', ''))
                                        keyCheck4 = (keyCheck3.replace(' ', ''))
                                        keyCheck4.split()
                                                                
                                        

                                        if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                            if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                                #report3.add_paragraph(dicts2Copy[str(keyCheck4)])
                                                print("")
                                            #orphanReport.add_paragraph(dicts2Copy[str(keyCheck4)])

                                        else:
                                            print("")
                                            #orphanChildren2Copy.append(str(keyCheck4))
                                            #report3.add_paragraph("Requirement text not found")
                                            #orphanReport.add_paragraph("Requirement text not found")
                                        #print(dicts10[str(key)])
                                        #report3.add_paragraph(dicts10[str(stringKey)])
                                        #for dicts10[str(stringKey)] in dicts10:
                                        #report3.add_paragraph(dicts10[str(stringKey)])
                                        for b in PTags:
                                            b = (b.replace(']', ''))
                                            #report3.add_paragraph("I'm b")
                                            #report3.add_paragraph(b)
                                            #report3.add_paragraph("I'm tag")
                                            #report3.add_paragraph(tag)


                                            if b == tag:
                                                i += 1
                                                hx = tag
                                                #report3.add_paragraph("I'm here")
                                                keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags
                                                #report3.add_paragraph("I'm keys")
                                                #report3.add_paragraph(keys)
                                                k += 1
                                                for item in keys: #keys are child tags of hx/the parent tag

                                                    if item != "" and item!= " ":
                                                        print("")
                                                        #report3.add_paragraph(item, style='List Bullet')
                                                        #para = report3.add_paragraph(dicts2Copy[str(item)])
                                                        #para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text





                        else:
                            print("")    
                            PTags = text.split(']')
                            PTags = [s.strip() + ']' for s in PTags]
                            PTags.pop()

                            for x in PTags:
                                keyCheck = (x.replace('[', ''))
                                keyCheck2 = (keyCheck.replace(']', ''))
                                keyCheck3 = (keyCheck2.replace(']', ''))
                                keyCheck4 = (keyCheck3.replace(' ', ''))
                                #report3.add_paragraph(x) # display the parent tag, included brackets

                                if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                    if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                        print("")
                                        #report3.add_paragraph(dicts2Copy[str(keyCheck4)])
                                    #orphanReport.add_paragraph(dicts2Copy[str(keyCheck4)])

                                else:
                                    #report3.add_paragraph("Requirement text not found")
                                    #orphanChildren2Copy.append(str(keyCheck4))
                                    print("")
                                    #orphanReport.add_paragraph("Requirement text not found")
                                #print(dicts10[str(key)])
                                #report3.add_paragraph(dicts10[str(stringKey)])
                                #for dicts10[str(stringKey)] in dicts10:
                                #report3.add_paragraph(dicts10[str(stringKey)])
                                for b in PTags:

                                    

                                    if b == dicts10[str(stringKey2)]:
                                        i += 1
                                        hx = dicts10[str(stringKey2)]
                                        keys = [h for h, v in dicts10.items() if v == hx] # finds all the child tags
                                        #print(keys)
                                        k += 1
                                        for item in keys: #keys are child tags of hx/the parent tag

                                            if item != "" and item!= " ":
                                                print("")
                                                #report3.add_paragraph(item, style='List Bullet')
                                                #para = report3.add_paragraph(dicts2Copy[str(item)])
                                                #para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text

                            #report3.add_paragraph("\n") # Adds a line space
                            #print(k)
                            #print(m)
                            #report3.add_paragraph(key, style='List Bullet')
                            #para = report3.add_paragraph(value)
                            #para.paragraph_format.left_indent = Inches(0.25) # adds indentation ot text
                            #stringKey = dicts2Copy[str(key)]
                            #stringKey2 = (stringKey.replace(' ', ''))
                        #if k < len(fullText2Copy):
                    #elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                
                        
                    elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                        k += 1
                        #orphanReport.add_paragraph("\n")
                        if i < len(parents2Copy):
                            orphanss.append(parents2Copy[i]) # adds orphan tags to a list
                            #orphanReport.add_paragraph(parents2Copy[i])
                            #orphanReport.add_paragraph(parents2Copy[i] + " is an orphan tag")
                            #print(parents2Copy[i])
                            #print(orphanTagText[o])
                            print("")
                        if o < len(orphanTagText):
                            #orphanReport.add_paragraph(orphanTagText[o])
                            print("")
                            
                        o += 1
                        if i < len(parents2Copy):
                            print("")
                            #orphanReport.add_paragraph(parents2Copy[i] + " is an orphan tag")
                        i += 1

        #orphanss.sort() # sorts the list of orphan tags
        #orphanChildren2Copy.sort() # sorts the list of orphan child tags
        """
        
        #for orph in orphanss:
        #    orphanReport.add_paragraph(orph)
        #orphanReport.add_paragraph("Orphan Tags2: ")


        for orph5 in orphanChildren2Copy:
            orphanReport.add_paragraph(orph5)
            print(orph5)

        orphanReport.save('orphanReport.docx')
        toggle_state4() # This will enable the open orphan report button
        return dicts2Copy

    except Exception as e:
        # Log an error message
        logging.error('orphanReport(): ERROR', e)
    else:
        # Log a success message
        logging.info('orphanReport(): PASS')



def removeParent(text): #removes parent tags or child tags
    try:
        childAfter = []
        for line in text:
            childAfter = [i.rsplit('[', 1)[0] for i in text] # removes parent tags
            childAfter = [re.sub("[\(\[].*?[\)\]]", "", e) for e in childAfter]  # removes parent tags that are left
            childAfter = [re.sub("[\{\[].*?[\)\}]", "", e) for e in childAfter]  # removes "pass", "fail", etc.
        return childAfter

    except Exception as e:
        # Log an error message
        logging.error('removeParent(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeParent(): PASS')


def removeText(text6): #this should remove everything before the parent tag
    try:
        childAfter = [s.split(None, 1)[0] if len(s.split(None, 1)) >= 2 else '' for s in text6]
        return childAfter
    except Exception as e:
        # Log an error message
        logging.error('removeText(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeText(): PASS')


def removeAfter(childtags): #removes everything after the  tag, example "pass"
    try:
        seperator = ']'
        childAfter = [i.rsplit(']', 1)[0] + seperator for i in childtags]

        return childAfter
    except Exception as e:
        # Log an error message
        logging.error('removeAfter(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeAfter(): PASS')


def removechild(text): #removes child, this one needs fixing
    try:
        mylst = []
        mylst = [s.split(None, 1)[1] if len(s.split(None, 1)) >= 2 else '' for s in text]
        return mylst
    except Exception as e:
        # Log an error message
        logging.error('removechild(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removechild(): PASS')

# This function will open up the report automatically
def getDocumentTable():
    try:
        if platform.system() == 'Darwin':
            subprocess.check_call(['open', 'ChildandParentTagsTables.docx'])
        elif platform.system() == 'Windows':
            os.startfile('ChildandParentTagsTables.docx')
        # os.startfile(report3) # try either one for windows if the first option gives error
        else:
            subprocess.call('xdg-open', report1)
    except Exception as e:
        # Log an error message
        logging.error('getDocumentTable(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('getDocumentTable(): PASS')


# This function will open up the report automatically
def getDocument():
    try:
        if platform.system() == 'Darwin':
            subprocess.check_call(['open', 'AllChildandParentTags.docx'])
        elif platform.system() == 'Windows':
            os.startfile('AllChildandParentTags.docx')
        # os.startfile(report3) # try either one for windows if the first option gives error
        else:
            subprocess.call('xdg-open', report3)
    except Exception as e:
        # Log an error message
        logging.error('getDocument(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('getDocument(): PASS')

def getOrphanDocument():
    try:
        if platform.system() == 'Darwin':
            subprocess.check_call(['open', 'orphanReport.docx'])
        elif platform.system() == 'Windows':
            os.startfile('orphanReport.docx')
        # os.startfile(orphanReport) # try either one for windows if the first option gives error
        else:
            subprocess.call('xdg-open', orphanGenReport)
    except Exception as e:
        # Log an error message
        logging.error('getOrphanDocument(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('getOrphanDocument(): PASS')


def getChildlessDocument():
    try:
        if platform.system() == 'Darwin':
            subprocess.check_call(['open', 'ChildlessTags.docx'])
        elif platform.system() == 'Windows':
            os.startfile('ChildlessTags.docx')
        # os.startfile(orphanReport) # try either one for windows if the first option gives error
        else:
            subprocess.call('xdg-open', generateReport2)
    except Exception as e:
        # Log an error message
        logging.error('getChildlessDocument(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('getChildlessDocument(): PASS')

def getTBV():
    try:
        if platform.system() == 'Darwin':
            subprocess.check_call(['open', 'TBVReport.docx'])
        elif platform.system() == 'Windows':
            os.startfile('TBVReport.docx')
        # os.startfile(tbvReport) # try either one for windows if the first option gives error
        else:
            subprocess.call('xdg-open', generateReport2)
    except Exception as e:
        # Log an error message
        logging.error('getTBV(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('getTBV(): PASS')
        
def getTBD():
    try:
        if platform.system() == 'Darwin':
            subprocess.check_call(['open', 'TBDReport.docx'])
        elif platform.system() == 'Windows':
            os.startfile('TBDReport.docx') 
        # os.startfile(tbvReport) # try either one for windows if the first option gives error
        else:
            subprocess.call('xdg-open', generateReport2)
    except Exception as e:
        # Log an error message
        logging.error('getTBD(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('getTBD(): PASS')

# Creates an excel report
def createExcel():
        
        wb = xw.Book()
        excelReport = wb.sheets[0]
        excelReport.name = "Report"

        # creating a Dataframe object from a list
        # of tuples of key, value pair
        df = pd.DataFrame(list(dicts2Copy.items()))
        # Dictionary For child and parent tag
        df2 = pd.DataFrame(list(dicts11111.items()))

        # For Orphan Tags
        df3 = pd.DataFrame(orphanChildren2Copy)

        # For Childless Tags
        df4 = pd.DataFrame(childless)

        df5 = pd.DataFrame(TBVTags)

        df6 = pd.DataFrame(TBDTags)
        

        # For childTag -Text
        excelReport.range("A1").value = df

        # Select the range with the dataframe
        #data_range = ws.range('A1').expand()
        # Drop the indexes
        #data_range.options(index=False).value

        # Listing out the Orphan Tags
        excelReport.range("G1").value = df3
        df3 = df.reset_index(drop=True)

        # Listing out the Childless Tags
        excelReport.range("I1").value = df4
        df4 = df.reset_index(drop=True)

        # Listing out the TBV Tags
        excelReport.range("K1").value = df5
        df5 = df.reset_index(drop=True)

        # Listing out the TBD Tags
        excelReport.range("M1").value = df6
        df6 = df.reset_index(drop=True)

        

        # Adding childTag header
        excelReport.range("B1").value = 'Child Tag'
        excelReport.range("B1").font.Size = 14 # Change font size
        excelReport.range("B1").font.ColorIndex = 2 # Change font color
        excelReport.range('B1:B1').color = (255, 0, 0) # Change cell background color

        # Adding Text header
        excelReport.range("C1").value = 'Text'
        excelReport.range("C1").font.Size = 14 # Change font size
        excelReport.range("C1").font.ColorIndex = 2 # Change font color
        excelReport.range('C1:C1').color = (0,255,0) # Change cell background color

        # For the childTag - parentTag
        excelReport.range("D1").value = df2

        excelReport.range("E1").value = "Child Tag"
        excelReport.range("E1").font.Size = 14
        excelReport.range("E1").font.ColorIndex = 2
        excelReport.range("E1:E1").color = (255, 0, 0)

        # Adding parentTag header
        excelReport.range("F1").value = 'Parent Tag'
        excelReport.range("F1").font.Size = 14 # Change font size
        excelReport.range("F1").font.ColorIndex = 2 # Change font color
        excelReport.range('F1:F1').color = (128, 128, 128) # Change cell background color

        # Adding OrphanTags header
        excelReport.range("H1").value = 'Orphan Tags'
        excelReport.range("H1").font.Size = 14 # Change font size
        excelReport.range("H1").font.ColorIndex = 2 # Change font color
        excelReport.range('H1:H1').color = (255, 128, 0) # Change cell background color

        # Adding ChildlessTags header
        excelReport.range("J1").value = 'Childless Tags'
        excelReport.range("J1").font.Size = 14 # Change font size
        excelReport.range("J1").font.ColorIndex = 2 # Change font color
        excelReport.range('J1:J1').color = (150, 75, 0) # Change cell background color

        # Adding TBVTags header
        excelReport.range("L1").value = 'TBV Tags'
        excelReport.range("L1").font.Size = 14 # Change font size
        excelReport.range("L1").font.ColorIndex = 2 # Change font color
        excelReport.range('L1:L1').color = (150, 75, 0) # Change cell background color

        # Adding TBDTags header
        excelReport.range("N1").value = 'TBD Tags'
        excelReport.range("N1").font.Size = 14 # Change font size
        excelReport.range("N1").font.ColorIndex = 2 # Change font color
        excelReport.range('N1:N1').color = (150, 75, 0) # Change cell background color
        
        
        excelReport.autofit()

        for key in dicts2:
            wb.sheets[0].append([key, dicts2[key]])
        wb.save('Tags&Requirements.xlsx') # Saving excel report as 'Tags&Requirements.xlsx'

"""
# Creates an excel report
def createExcel2():
    try:
        # book_arr = xw.App().books.add()
        # wb = book_arr.add()
        # wb.title = "Report"
        
        
        
        # excelReport = wb.sheets.add("Report")

        #excelReport.name = report
        excelReport2.range("B1").value = "Children"
        excelReport2.range("B1").font.Size = 18 # Change font size
        excelReport2.range("B1").font.ColorIndex = 2 # Change font color
        excelReport2.range('B1:B1').color = (255, 0, 0) # Change cell background color

        excelReport2.range("A1").value = 'Parents'
        excelReport2.range("A1").font.Size = 18 # Change font size
        excelReport2.range("A1").font.ColorIndex = 2 # Change font color
        excelReport2.range('A1:A1').color = (0, 0, 255) # Change cell background color



        # For Orphan Tags
        df3 = pd.DataFrame(orphanChildren2Copy)

        # For Childless Tags
        df4 = pd.DataFrame(childless)

        # For TBV Tags
        df5 = pd.DataFrame(TBVTags)

        # For TBD Tags
        df6 = pd.DataFrame(TBDTags)
        
        


        # Select the range with the dataframe
        #data_range = ws.range('A1').expand()
        # Drop the indexes
        #data_range.options(index=False).value

        # Listing out the Orphan Tags
        excelReport2.range("H1").value = df3
        df3 = df3.reset_index(drop=True)

        # Listing out the Childless Tags
        excelReport2.range("K1").value = df4
        df4 = df4.reset_index(drop=True)

        # Listing out the TBV Tags
        excelReport2.range("N1").value = df5
        df5 = df5.reset_index(drop=True)

        # Listing out the TBD Tags
        excelReport2.range("Q1").value = df6
        df6 = df6.reset_index(drop=True)
        

        
        # Adding OrphanTags header
        excelReport2.range("I1").value = 'Orphan Tags'
        excelReport2.range("I1").font.Size = 14 # Change font size
        excelReport2.range("I1").font.ColorIndex = 2 # Change font color
        excelReport2.range('I1:I1').color = (255, 128, 0) # Change cell background color

        # Adding ChildlessTags header
        excelReport2.range("L1").value = 'Childless Tags'
        excelReport2.range("L1").font.Size = 14 # Change font size
        excelReport2.range("L1").font.ColorIndex = 2 # Change font color
        excelReport2.range('L1:L1').color = (150, 75, 0) # Change cell background color

        # Adding TBVTags header
        excelReport2.range("O1").value = 'TBV Tags'
        excelReport2.range("O1").font.Size = 14 # Change font size
        excelReport2.range("O1").font.ColorIndex = 2 # Change font color
        excelReport2.range('O1:O1').color = (150, 75, 0) # Change cell background color

        # Adding TBDTags header
        excelReport2.range("R1").value = 'TBD Tags'
        excelReport2.range("R1").font.Size = 14 # Change font size
        excelReport2.range("R1").font.ColorIndex = 2 # Change font color
        excelReport2.range('R1:R1').color = (150, 75, 0) # Change cell background color
        

        
        
        excelReport2.autofit()



        wb2.save('AllTags.xlsx') # Saving excel report as 'AllTags.xlsx'
    except Exception as e:
        # Log an error message
        logging.error('createExcel2(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('createExcel2(): PASS')
"""

def toggle_state(): # this will re-enable getDoc button
    Gui.getDoc.config(state="normal")

def toggle_state2(): # this will re-enable generate report button
    Gui.genRep.config(state="normal")

def toggle_state3(): # this will re-enable excel report button
    Gui.getExcel.config(state="normal")

def toggle_state4(): # this will re-enable word report button for orphan tags
    Gui.getOrphanDoc.config(state="normal")

#def toggle_state5(): # this will re-enable excel report button for orphan tags
 #   Gui.getOrphan.config(state="normal")

def toggle_state6(): # this will re-enable allTags report button for tables
    Gui.allTagsButton.config(state="normal")

def toggle_state7(): # this will re-enable childless report button
    Gui.getChildlessDoc.config(state="normal")

def toggle_state8(): # this will re-enable excelreport 2 button
    Gui.getExcel2.config(state="normal")

def toggle_state9(): # this will re-enable tbv report button
    Gui.getTBVdoc.config(state="normal")

def toggle_state10(): # this will renable tbd report button
    Gui.getTBDdoc.config(state="normal")

def toggle_state13(): # this will renable treeview button
    Gui.TreeDiagram.config(state="normal")

#def check_string(string1, string2): # checks if a string1 is in string2
 #   if isinstance(string2, str): 
  #      string2 = [string2]
   # pattern = r'{}(?!\d)'.format(re.escape(string1))
    #for s in string2:
     #   match = re.search(pattern, s)
      #  if match is not None:
       #     return True
    #return False



def check_string2(string1, string2):
    if not isinstance(string1, str):
        return False
    if isinstance(string2, str):
        string2 = [string2]
    elif not isinstance(string2, list):
        return False
    
    pattern = r'{}(?!\d)'.format(re.escape(string1))
    
    for s in string2:
        if not isinstance(s, str):
            continue
        match = re.search(pattern, s)
        if match is not None:
            return True
    
    return False


def check_string(string1, string2):
    if not isinstance(string1, str):
        return False
    if isinstance(string2, str):
        string2 = [string2]
    elif not isinstance(string2, list):
        return False
    
    pattern = r'(?<=\[)?{}(?=\])?(?!\d)'.format(re.escape(string1))
    
    for s in string2:
        if not isinstance(s, str):
            continue
        match = re.search(pattern, s)
        if match is not None:
            return True
    
    return False




import platform
try:
    import win32com.client
    win32com_available = True
except ModuleNotFoundError:
    win32com_available = False
import subprocess

def closeReports():
    try:
        if platform.system() == 'Windows' and win32com_available:
            word = win32com.client.Dispatch("Word.Application")
            for doc in word.Documents:
                doc.Close()
            word.Quit()
        elif platform.system() == 'Darwin':
            # Use subprocess to control Microsoft Word on macOS
            subprocess.call(['osascript', '-e', 'tell application "Microsoft Word" to quit'])
    except Exception as e:
        # Log an error message
        logging.error('closeDocuments(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('closeDocuments(): PASS')

def closeExcelWorkbooks():
    try:
        if platform.system() == 'Windows' and win32com_available:
            excel = win32com.client.Dispatch("Excel.Application")
            for book in excel.Workbooks:
                book.Close(SaveChanges=False)
            excel.Quit()
        elif platform.system() == 'Darwin':
            # Use subprocess to control Microsoft Excel on macOS
            subprocess.call(['osascript', '-e', 'tell application "Microsoft Excel" to quit'])
    except Exception as e:
        # Log an error message
        logging.error('closeExcelWorkbooks(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('closeExcelWorkbooks(): PASS')


def remove_empty_values(d):
    """
    This function takes a dictionary 'd' as input and removes all key-value pairs from it where the value is an empty string, whitespace, or an empty list.
    """
    keys_to_remove = []
    for key, value in d.items():
        if isinstance(value, list):
            if not value:
                keys_to_remove.append(key)
            else:
                for v in value:
                    if v.strip() == "":
                        keys_to_remove.append(key)
                        break
        elif value.strip() == "":
            keys_to_remove.append(key)
    for key in keys_to_remove:
        del d[key]



def createExcel3():
    
    global wb3
    wb3 = xw.Book()
    excelReport3 = wb3.sheets[0]
    excelReport3.name = "ReportNew"

    # Adding First Generation Tags header
    excelReport3.range("A1").value = 'First Generation'
    excelReport3.range("A1").font.Size = 14 # Change font size
    excelReport3.range("A1").font.ColorIndex = 2 # Change font color
    excelReport3.range('A1').font.bold = True
    excelReport3.range('A1:A1').color = (121,205,205) # Change cell background color

    # Adding Second Generation Tags header
    excelReport3.range("B1").value = 'Second Generation'
    excelReport3.range("B1").font.Size = 14 # Change font size
    excelReport3.range('B1').font.bold = True
    excelReport3.range("B1").font.ColorIndex = 2 # Change font color
    excelReport3.range('B1:B1').color = (121,205,205) # Change cell background color

    # Adding Third Generation Tags header
    excelReport3.range("C1").value = 'Third Generation'
    excelReport3.range("C1").font.Size = 14 # Change font size
    excelReport3.range('C1').font.bold = True
    excelReport3.range("C1").font.ColorIndex = 2 # Change font color
    excelReport3.range('C1:C1').color = (121,205,205) # Change cell background color

    # Adding Fourth Generation Tags header
    excelReport3.range("D1").value = 'Fourth Generation'
    excelReport3.range("D1").font.Size = 14 # Change font size
    excelReport3.range('D1').font.bold = True
    excelReport3.range("D1").font.ColorIndex = 2 # Change font color
    excelReport3.range('D1:D1').color = (121,205,205) # Change cell background color

    # Adding Fifth Generation Tags header
    excelReport3.range("E1").value = 'Fifth Generation'
    excelReport3.range("E1").font.Size = 14 # Change font size
    excelReport3.range('E1').font.bold = True
    excelReport3.range("E1").font.ColorIndex = 2 # Change font color
    excelReport3.range('E1:E1').color = (121,205,205) # Change cell background color

    # Adding Sixth Generation Tags header
    excelReport3.range("F1").value = 'Sixth Generation'
    excelReport3.range("F1").font.Size = 14 # Change font size
    excelReport3.range('F1').font.bold = True
    excelReport3.range("F1").font.ColorIndex = 2 # Change font color
    excelReport3.range('F1:F1').color = (121,205,205) # Change cell background color

    # Adding Fifth Generation Tags header
    excelReport3.range("G1").value = 'Seventh Generation'
    excelReport3.range("G1").font.Size = 14 # Change font size
    excelReport3.range('G1').font.bold = True
    excelReport3.range("G1").font.ColorIndex = 2 # Change font color
    excelReport3.range('G1:G1').color = (121,205,205) # Change cell background color

    # counters for Excel report3

    global counter1
    counter1 = 2
    global counter2
    counter2 = 1
    global counter3
    counter3 = 0
    global cell
    cell = 0;
    global cell2
    cell2 = 0;
    #print("dicts10",dicts10)
    remove_empty_values(dicts10)
    #print("new dicts10",dicts10)
    #print("this is dicts10",dicts10)

    hello = 'PUMP:RISK:10'


    for orpha in orphanChildren2Copy:
        if orpha != orphanChildren2Copy[0]:
            cell10 = str('A'+ str(counter1-1))
            excelReport3.range(cell10).value = 'SEPARATOR'
            excelReport3.range(cell10).font.Size = 14 # Change font size
            excelReport3.range(cell10).font.ColorIndex = 2 # Change font color
            cell11 = str(str(cell10) + ':G' + str(counter1-1))
            excelReport3.range(cell11).color = (178,223,238) # Change cell background color

        cell = str('A'+ str(counter1))
        cell2 = str(str(orpha))
        excelReport3.range(cell).value = cell2
        counter1 += 2 # counter for excel report
        counter2 += 1 # counter for excel report
        #print("hello",len(hello))
        #print(type(orpha))
        hx = str(orpha)
        #hx.replace(" ", "")
        #print(repr(hx))
        #print(hx)
        hx = hx.strip('\n')  # remove newline characters
        hx = hx.strip()     # remove leading/trailing whitespaces
        #print(repr(hx))
        #print(hx)
        #print(type(hx))
        hx.strip()
        #print(len(hx))
        keys = [h for h, v in dicts10.items() if check_string(hx, str(v))]
        #wi = 'PUMP:RISK:10'
        #keys3 = [h for h, v in dicts10.items() if check_string(wi, str(v))] # finds all the child tags
        #print("keys3",keys3)

        #print("keys",keys)
        for item in keys: #keys are child tags of hx/the parent tag
             #print("child",item)
             if item != "" and item!= " ": # if the child tag is not empty

                counter2 = counter1 - 1
                cell = str('B'+ str(counter2))                                                
                cell2 = str(item)  
                #print("cell",cell) 
                #print("cell2",cell2)                                              
                excelReport3.range(cell).value = cell2
                counter2 += 1
                counter1 += 1
                stringKey = str(item)
                stringKey2 = (stringKey.replace(' ', ''))
                #duplicates2 = []
                #print("item is not '' ", item)
                if str(stringKey2) in dicts10: # if the key is in the dictionary
                #for key in dicts10:
                    if item in dicts10:
                        #print("item is in dicts10",item)
                        hx = item

                        keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags

                        for item1 in keys: #keys are child tags of hx/the parent tag
                            #print("grandchild",item1)
                            if item1 != "" and item1!= " ":
                                #if item1 not in duplicates2:
                                    #duplicates2.append(item1)
                                    counter2 = counter1 - 1
                                    cell = str('C'+ str(counter2))                                                
                                    cell2 = str(item1)                                                
                                    excelReport3.range(cell).value = cell2
                                    counter2 += 1
                                    counter1 += 1
                                    stringKey = str(item1)
                                    stringKey2 = (stringKey.replace(' ', ''))
                                    #duplicates2 = []
                                    #print("item is not '' ", item)
                                    if str(stringKey2) in dicts10: # if the key is in the dictionary
                                    #for key in dicts10:
                                        if item1 in dicts10:
                                            #print("item is in dicts10",item)
                                            hx = item1

                                            keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags

                                            for item2 in keys: #keys are child tags of hx/the parent tag
                                                #print("grandgrandchild",item2)
                                                if item2 != "" and item2!= " ":
                                                    #if item2 not in duplicates2:
                                                        #duplicates2.append(item1)
                                                        counter2 = counter1 - 1
                                                        cell = str('D'+ str(counter2))                                                
                                                        cell2 = str(item2)                                                
                                                        excelReport3.range(cell).value = cell2
                                                        counter2 += 1
                                                        counter1 += 1
                                                        stringKey = str(item2)
                                                        stringKey2 = (stringKey.replace(' ', ''))
                                                        #duplicates2 = []
                                                        #print("item is not '' ", item)
                                                        if str(stringKey2) in dicts10: # if the key is in the dictionary
                                                        #for key in dicts10:
                                                            if item2 in dicts10:
                                                                #print("item is in dicts10",item)
                                                                hx = item2

                                                                keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags

                                                                for item3 in keys: #keys are child tags of hx/the parent tag
                                                                    #print("grandgrandGrandchild",item3)
                                                                    if item3 != "" and item3!= " ":
                                                                        #if item2 not in duplicates2:
                                                                            #duplicates2.append(item1)
                                                                            counter2 = counter1 - 1
                                                                            cell = str('E'+ str(counter2))                                                
                                                                            cell2 = str(item3)                                                
                                                                            excelReport3.range(cell).value = cell2
                                                                            counter2 += 1
                                                                            counter1 += 1
                                                                            stringKey = str(item3)
                                                                            stringKey2 = (stringKey.replace(' ', ''))
                                                                            #duplicates2 = []
                                                                            #print("item is not '' ", item)
                                                                            if str(stringKey2) in dicts10: # if the key is in the dictionary
                                                                            #for key in dicts10:
                                                                                if item3 in dicts10:
                                                                                    #print("item is in dicts10",item)
                                                                                    hx = item3

                                                                                    keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags

                                                                                    for item4 in keys: #keys are child tags of hx/the parent tag
                                                                                        #print("grandgrandGrandGrandchild",item3)
                                                                                        if item4 != "" and item4!= " ":
                                                                                            #if item2 not in duplicates2:
                                                                                                #duplicates2.append(item1)
                                                                                                counter2 = counter1 - 1
                                                                                                cell = str('F'+ str(counter2))                                                
                                                                                                cell2 = str(item4)                                                
                                                                                                excelReport3.range(cell).value = cell2
                                                                                                counter2 += 1
                                                                                                counter1 += 1
                                                                                                stringKey = str(item4)
                                                                                                stringKey2 = (stringKey.replace(' ', ''))
                                                                                                #duplicates2 = []
                                                                                                #print("item is not '' ", item)
                                                                                                if str(stringKey2) in dicts10: # if the key is in the dictionary
                                                                                                #for key in dicts10:
                                                                                                    if item4 in dicts10:
                                                                                                        #print("item is in dicts10",item)
                                                                                                        hx = item4

                                                                                                        keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags

                                                                                                        for item5 in keys: #keys are child tags of hx/the parent tag
                                                                                                            #print("grandgrandGrandGrandchild",item3)
                                                                                                            if item5 != "" and item5!= " ":
                                                                                                                #if item2 not in duplicates2:
                                                                                                                    #duplicates2.append(item1)
                                                                                                                    counter2 = counter1 - 1
                                                                                                                    cell = str('G'+ str(counter2))                                                
                                                                                                                    cell2 = str(item5)                                                
                                                                                                                    excelReport3.range(cell).value = cell2
                                                                                                                    counter2 += 1
                                                                                                                    counter1 += 1
                                                                                                                    stringKey = str(item5)
                                                                                                                    stringKey2 = (stringKey.replace(' ', ''))
                                                                                                                    #duplicates2 = []
                                                                                                                    #print("item is not '' ", item)
                                                                                                                    if str(stringKey2) in dicts10: # if the key is in the dictionary
                                                                                                                    #for key in dicts10:
                                                                                                                        if item5 in dicts10:
                                                                                                                            #print("item is in dicts10",item)
                                                                                                                            hx = item5

                                                                                                                            keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags

                                                                                                                            for item6 in keys: #keys are child tags of hx/the parent tag
                                                                                                                                #print("grandgrandGrandGrandchild",item3)
                                                                                                                                if item6 != "" and item6!= " ":
                                                                                                                                    #if item2 not in duplicates2:
                                                                                                                                        #duplicates2.append(item1)
                                                                                                                                        counter2 = counter1 - 1
                                                                                                                                        cell = str('H'+ str(counter2))                                                
                                                                                                                                        cell2 = str(item6)                                                
                                                                                                                                        excelReport3.range(cell).value = cell2
                                                                                                                                        counter2 += 1
                                                                                                                                        counter1 += 1



    #dicts9999 = {'PUMP:HRD:100': '[PUMP:HRS:100]', 'PUMP:HRD:105': '[PUMP:HRS:103]', 'PUMP:HRD:1000': '[PUMP:HRS:1000]', 'PUMP:HRD:3330': '[PUMP:HRS:3330]', 'PUMP:HRD:3350': '[PUMP:TBV:1111]', 'PUMP:HRD:0000': '[PUMP:TBV:1111]', 'PUMP:HRS:100': '[PUMP:PRS:100]', 'PUMP:HRS:105': '[PUMP:PRS:103]', 'PUMP:HRS:1000': '[PUMP:PRS:1000]', 'PUMP:HRS:3330': '[PUMP:PRS:3330]', 'PUMP:HRS:3340': '[PUMP:PRS:3330]', 'PUMP:HRS:3350': '[PUMP:PRS:3350]', 'PUMP:HTP:100': '[PUMP:HRS:100]', 'PUMP:HTP:200': '[PUMP:HRS:105]', 'PUMP:HTP:300': '[PUMP:HRS:1000]', 'PUMP:HTP:400': '[PUMP:HRS:3330]', 'PUMP:HTP:500': '[PUMP:HRS:3350]', 'PUMP:HTP:1100': '[PUMP:HRD:100]', 'PUMP:HTP:1200': '[PUMP:HRD:105]', 'PUMP:HTP:1300': '[PUMP:HRD:1000]', 'PUMP:HTP:1400': '[PUMP:HRD:3330]', 'PUMP:HTP:1500': '[PUMP:HRD:3350]', 'PUMP:HTR:100': '[PUMP:HTP:100]', 'PUMPHTR:200': '[PUMP:HTP:200]', 'PUMP:HTR:300': '[PUMP:HTP:300]', 'PUMP:HTR:400': '[PUMP:HTP:400]', 'PUMP:HTR:500': '[PUMP:HTP:500]', 'PUMP:HTR:1100': '[PUMP:HTP:1100]', 'PUMP:HTR:1200': '[PUMP:HTP:1200]', 'PUMP:HTR:1300': '[PUMP:HTP:1300]', 'PUMP:HTR:1400': '[PUMP:HTP:1400]', 'PUMP:HTR:1500': '[PUMP:HTP:1500]', 'PUMP:PRS:1': '[PUMP:URS:1]', 'PUMP:PRS:2': '[PUMP:RISK:10]', 'PUMP:PRS:3': '[PUMP:RISK:20]', 'PUMP:PRS:4': '[PUMP:URS:3]', 'PUMP:PRS:5': '[PUMP:URS:3]', 'PUMP:PRS:8': ['PUMP:URS:8', 'PUMP:RISK:30'], 'PUMP:PRS:10': '[PUMP:URS:10]', 'PUMP:PRS:100': '[PUMP:URS:100]', 'PUMP:PRS:105': ['PUMP:URS:103', 'PUMP:RISK:40'], 'PUMP:PRS:1000': ['PUMP:URS:1000', 'PUMP:RISK:50'], 'PUMP:PRS:3330': '[PUMP:URS:3330]', 'PUMP:PRS:3340': '[PUMP:URS:3330]', 'PUMP:PRS:3350': '[PUMP:URS:3350]', 'PUMP:PRS:4000': '[PUMP:URS:4000]', 'PUMP:RISK:10': '', 'PUMP:RISK:20': '', 'PUMP:RISK:30': '', 'PUMP:RISK:40': '', 'PUMP:RISK:50': '', 'PUMP:SDS:10': ['BOLUS:SRS:1', 'BOLUS:SRS:2', 'BOLUS:SRS:5', 'BOLUS:SRS:6', 'BOLUS:SRS:8', 'BOLUS:SRS:12', 'ACE:SRS:1', 'ACE:SRS:5', 'ACE:SRS:6'], 'PUMP:SDS:20': '[ACE:SRS:2]', 'PUMP:SDS:30': ['AID:SRS:1', 'AID:SRS:2', 'AID:SRS:10', 'AID:SRS:12', 'AID:SRS:20'], 'PUMP:SDS:40': ['ACE:SRS:110', 'ACE:SRS:120'], 'PUMP:SDS:50': '[ACE:SRS:110]', 'PUMP:SDS:60': '[ACE:SRS:10]', 'PUMP:SDS:70': '[ACE:SRS:100]', 'ACE:SRS:1': ['PUMP:PRS:1', 'PUMP:TBV:1'], 'ACE:SRS:2': '[PUMP:PRS:1]', 'ACE:SRS:5': '[PUMP:PRS:5]', 'ACE:SRS:6': '[PUMP:PRS:6]', 'ACE:SRS:10': '[PUMP:PRS:10]', 'ACE:SRS:100': '[PUMP:PRS:105]', 'BOLUS:SRS:1': '[PUMP:PRS:1]', 'BOLUS:SRS:2': ['PUMP:PRS:1', 'PUMP:TBD:1'], 'BOLUS:SRS:5': ['PUMP:PRS:1', 'PUMP:PRS:5'], 'BOLUS:SRS:6': ['PUMP:PRS:1', 'PUMP:PRS:3'], 'BOLUS:SRS:8': '[PUMP:PRS:1]', 'BOLUS:SRS:12': ['PUMP:PRS:1', 'PUMP:PRS:8'], 'AID:SRS:1': ['PUMP:PRS:4000', 'PUMP:DER:2'], 'AID:SRS:2': ['PUMP:PRS:4000', 'PUMP:DER:2'], 'AID:SRS:10': ['PUMP:PRS:4000', 'PUMP:DER:2'], 'AID:SRS:12': ['PUMP:PRS:4000', 'PUMP:DER:2'], 'AID:SRS:20': ['PUMP:PRS:4000', 'PUMP:DER:2'], 'PUMP:SVAL:100': ['ACE:SRS:1', 'ACE:SRS:5', 'ACE:SRS:6', 'BOLUS:SRS:1', 'BOLUS:SRS:2', 'BOLUS:SRS:5', 'BOLUS:SRS:6', 'BOLUS:SRS:8'], 'PUMP:SVAL:200': '[ACE:SRS:2]', 'PUMP:SVAL:300': ['ACE:SRS:10', 'ACE:SRS:100', 'ACE:SRS:1000', 'ACE:SRS:120'], 'PUMP:SVAL:400': '[BOLUS:SRS:12]', 'PUMP:SVAL:500': ['AID:SRS:1', 'AID:SRS:2', 'AID:SRS:10', 'AID:SRS:12', 'AID:SRS:20'], 'PUMP:SVATR:100': '[PUMP:SVAL:100]', 'PUMP:SVATR:200': '[PUMP:SVAL:200]', 'PUMP:SVATR:300': '[PUMP:SVAL:300]', 'PUMP:SVATR:400': '[PUMP:SVAL:400]', 'PUMP:SVATR:500': '[PUMP:SVAL:500]', 'PUMP:UT:100': '[PUMP:UNIT:100]', 'PUMP:UT:110': '[PUMP:UNIT:110]', 'PUMP:UT:120': '[PUMP:UNIT:120]', 'PUMP:UT:130': '[PUMP:UNIT:130]', 'PUMP:UT:140': '[PUMP:UNIT:140]', 'PUMP:UT:150': '[PUMP:UNIT:150]', 'PUMP:UT:160': '[PUMP:UNIT:160]', 'PUMP:UT:170': '[PUMP:UNIT:170]', 'PUMP:UT:180': '[PUMP:UNIT:180]', 'PUMP:UT:190': '[PUMP:UNIT:190]', 'PUMP:UT:200': '[PUMP:UNIT:200]', 'PUMP:UT:210': '[PUMP:UNIT:210]', 'PUMP:UT:220': '[PUMP:UNIT:220]', 'PUMP:INS:100': '[PUMP:UNIT:100]', 'PUMP:INS:110': '[PUMP:UNIT:110]', 'PUMP:INS:120': '[PUMP:UNIT:120]', 'PUMP:INS:130': '[PUMP:UNIT:130]', 'PUMP:INS:140': '[PUMP:UNIT:140]', 'PUMP:INS:150': '[PUMP:UNIT:150]', 'PUMP:INS:160': '[PUMP:UNIT:160]', 'PUMP:INS:170': '[PUMP:UNIT:170]', 'PUMP:INS:180': '[PUMP:UNIT:180]', 'PUMP:INS:190': '[PUMP:UNIT:190]', 'PUMP:INS:200': '[PUMP:UNIT:200]', 'PUMP:INS:210': '[PUMP:UNIT:210]', 'PUMP:INS:220': '[PUMP:UNIT:220]', 'PUMP:URS:1': '', 'PUMP:URS:3': '', 'PUMP:URS:8': '', 'PUMP:URS:10': '', 'PUMP:URS:100': '', 'PUMP:URS:103': '', 'PUMP:URS:1000': '', 'PUMP:URS:3330': '', 'PUMP:URS:3350': '', 'PUMP:URS:4000': ''}
    #hx = 'PUMP:RISK:10'
    #print(type(hx))
    #keys = [h for h, v in dicts9999.items() if check_string(hx, v)]
    #print("parent",hx)
    #print("keys2", keys)
    excelReport3.autofit()
    wb3.save('excelNew.xlsx') # Saving excel report as 'AllTags.xlsx'


def guiTree():
    family_trees = []
    current_tree = []

    for orpha in orphanChildren2Copy:
        if orpha != orphanChildren2Copy[0]:
            family_trees.append(current_tree)
            current_tree = []

        hx = orpha.strip()
        current_tree += "├─ " + hx + "\n"

        keys = [h for h, v in dicts10.items() if check_string(hx, str(v))]

        for i, item in enumerate(keys): #keys are child tags of hx/the parent tag
            if item != "" and item != " ":
                current_tree += "│  ├─ " + item + "\n"
                stringKey = item.strip()
                if stringKey in dicts10:
                    hx = item
                    keys1 = [h for h, v in dicts10.items() if check_string(hx, v)]
                    for j, item1 in enumerate(keys1):
                        if item1 != "" and item1 != " ":
                            current_tree += "│  │  ├─ " + item1 + "\n"
                            stringKey = item1.strip()
                            if stringKey in dicts10:
                                hx = item1
                                keys2 = [h for h, v in dicts10.items() if check_string(hx, v)]
                                for k, item2 in enumerate(keys2):
                                    if item2 != "" and item2 != " ":
                                        current_tree += "│  │  │  ├─ " + item2 + "\n"
                                        stringKey = item2.strip()
                                        if stringKey in dicts10:
                                            hx = item2
                                            keys3 = [h for h, v in dicts10.items() if check_string(hx, v)]
                                            for l, item3 in enumerate(keys3):
                                                if item3 != "" and item3 != " ":
                                                    current_tree += "│  │  │  │  ├─ " + item3 + "\n"
                                                    stringKey = item3.strip()
                                                    if stringKey in dicts10:
                                                        hx = item3
                                                        keys4 = [h for h, v in dicts10.items() if check_string(hx, v)]
                                                        for m, item4 in enumerate(keys4):
                                                            if item4 != "" and item4 != " ":
                                                                current_tree += "│  │  │  │  │  ├─ " + item4 + "\n"
                                                                stringKey = item4.strip()
                                                                if stringKey in dicts10:
                                                                    hx = item4
                                                                    keys5 = [h for h, v in dicts10.items() if check_string(hx, v)]
                                                                    for n, item5 in enumerate(keys5):
                                                                        if item5 != "" and item5 != " ":
                                                                            current_tree += "│  │  │  │  │  │  ├─ " + item5 + "\n"
                                                                            stringKey = item5.strip()
                                                                            if stringKey in dicts10:
                                                                                hx = item5
                                                                                keys6 = [h for h, v in dicts10.items() if check_string(hx, v)]
                                                                                for o, item6 in enumerate(keys6):
                                                                                    if item6 != "" and item6 != " ":
                                                                                        current_tree += "│  │  │  │  │  │  │  ├─ " + item6 + "\n"
                                            

    family_trees.append(current_tree)
    return family_trees

            




                                                
                                                   


    
                                                            
                                                    
                                                    
                                                    

