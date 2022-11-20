# Importing required widgets
# from PyQt5.QtGui import *
# from PyQt5.QtCore import *
# from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QApplication, QListView, QAbstractItemView, QTreeView, QLabel, QPushButton
from PyQt5.QtWidgets import QWidget, QComboBox, QFileDialog, QVBoxLayout
import re
import sys
import os
import platform
import subprocess
import copy
import docx
from docx import Document
from docx.shared import RGBColor
from typing import Tuple


class MyApp(QWidget):
    def __init__(self):
        super().__init__()
        self.window_width, self.window_height = 200, 240
        self.setMinimumSize(self.window_width, self.window_height)

        layout = QVBoxLayout()
        self.setLayout(layout)

        #self.options = ('Choose Document', 'Choose Documents','Generate Single Report', 'Choose Folder', 'Save File', 'Open Document')

        # self.combo = QComboBox()
        # self.combo.addItems(self.options)
        # layout.addWidget(self.combo)

        openFile = QPushButton('Choose Document')
        openFile.clicked.connect(self.openFile)
        layout.addWidget(openFile)

        openFiles = QPushButton('Choose Documents')
        openFiles.clicked.connect(self.getFileNames)
        layout.addWidget(openFiles)

        generateReport = QPushButton('Generate Single Report')
        generateReport.clicked.connect(self.generateReport)
        layout.addWidget(generateReport)

        openFolder = QPushButton('Choose Folder')
        openFolder.clicked.connect(self.getDirectory)
        layout.addWidget(openFolder)

        saveFile = QPushButton('Save File')
        saveFile.clicked.connect(self.getSaveFileName)
        layout.addWidget(saveFile)

        getDoc = QPushButton('Open Document')
        getDoc.clicked.connect(self.getDocument)
       # layout.addWidget(getDoc)

    # reads the text in the document and use the getcoloredTXT function
    def readtxt(filename, color: Tuple[int, int, int]):
        doc = docx.Document(filename)
        text10 = ""
        fullText = []
        new = []
        global everything
        everything = []  # list of tags and text

        for para in doc.paragraphs:
            # Getting the colored words from the doc
            if (filename.getcoloredTxt(para.runs, color)):
                # Concatenating list of runs between the colored text to single a string
                sentence = "".join(r.text for r in para.runs)
                fullText.append(sentence)
                # print(sentence) # Prints everything in the terminal
                everything.append(sentence)
                text10 = sentence
                parent.append("".join(r.text for r in para.runs))

        # print(fullText)
        global filtered_L  # Will store the ones without a child tag
        global hasChild  # Will store the ones with a child tag
        global fullText2  # will store everything found
        global children
        # Finds the lines without a childTag
        filtered_L = [value for value in fullText if "[" not in value]
        # Finds the lines with a childTag
        hasChild = [value for value in fullText if "[" in value]
        # will store everything found
        fullText2 = [value for value in fullText]

        return fullText, filtered_L, hasChild

    def getcoloredTxt(runs, color):  # Will look for colored text

        coloredWords, word = [], ""
        for run in runs:
            if run.font.color.rgb == RGBColor(*color):
                word += str(run.text)  # Saves everything found


            elif word != "":  # This will find the parentTags
                coloredWords.append(word)
                parentTags.append(word)
                parents.append(word)
                word = ""

        if word != "":  # This will find the parentTags
            coloredWords.append(word + "\n")
            # word = removeAfter(word)
            child.append(word)
            withChild.append(word)

        return coloredWords  # returns everything found

    # def launchDialog(self):
    #     option = self.options.index(self.combo.currentText())
    #     if option == 0:
    #         response = self.openFile()
    #     else:
    #         print('Nothing was Chosen')

    # def launchDialog1(self):
    #     option = self.options.index(self.combo.currentText())
    #     if option == 1:
    #         response = self.getFileNames()
    #     else:
    #         print('Nothing was Chosen')

    # def launchDialog2(self):
    #     option = self.options.index(self.combo.currentText())
    #     if option == 2:
    #         response = self.generateReport()
    #     else:
    #         print('Nothing was Chosen')

    # def launchDialog3(self):
    #     option = self.options.index(self.combo.currentText())
    #     if option == 3:
    #         response = self.getDirectory()
    #     else:
    #         print('Nothing was Chosen')

    # def launchDialog4(self):
    #     option = self.options.index(self.combo.currentText())
    #     if option == 4:
    #         response = self.getSaveFileName()
    #     else:
    #         print('Nothing was Chosen')

    # def launchDialog5(self):
    #     option = self.options.index(self.combo.currentText())
    #     if option == 5:
    #         response = self.getDocument()
    #     else:
    #         print('Nothing was Chosen')

    def openFile(self):
        global response
        global response2
        fileFilter = 'Word Document (*.docx);; All File Types (*.*)'
        response = QFileDialog.getOpenFileName(parent=self,
                                               caption='Select a File',
                                               directory=os.getcwd(),
                                               filter=fileFilter,
                                            )
        response2 = str(response)
        file = open(response2, 'r')
        #print(response)
        file.close()
        return response2[0]

    def getFileNames(self):
        fileFilter = 'Word Document (*.docx);; All File Types (*.*)'
        response = QFileDialog.getOpenFileNames(parent=self,
                                                caption='Select a File',
                                                directory=os.getcwd(),
                                                filter=fileFilter,
                                                )
        return response[0]

    def generateReport(self):  # Will generate the report for tags
        fullText = self.readtxt(filename=response2,
                                color=(255, 0, 0))
        # filtered_L = readtxt(filename=filepath2, #For future use
        #                   color=(255, 0, 0))
        fullText10 = str(fullText)
        s = ''.join(fullText10)
        w = (s.replace(']', ']\n\n'))
        paragraph = report3.add_paragraph()
        filepath3 = str(response2.rsplit('/', 1)[-1])  # change filepath to something.docx
        filepath3 = filepath3.split('.', 1)[0]  # removes .docx of the file name
        print(filepath3 + " report created")
        runner = paragraph.add_run("\n" + "Document Name: " + filepath3 + "\n")
        runner.bold = True  # makes the header bold
        # w will be used in the future
        w = (w.replace('([', ''))
        w = (w.replace(',', ''))
        w = (w.replace('' '', ''))

        # creates a  table
        table = report3.add_table(rows=1, cols=2)

        # Adds headers in the 1st row of the table
        row = table.rows[0].cells
        row[0].text = 'Child Tag'
        row[1].text = 'Parent Tag/tags'
        # Adding style to a table
        table.style = 'Colorful List'

        # Now save the document to a location
        report3.save('report3.docx')
        e = 0

        child2 = self.removeAfter(child)  # removes everything after the child tag if there is anything to remove
        # while loop until all the  parentTags has been added to the report

        parents2 = parentTags  # copy of parent tags list
        childCopy = copy.deepcopy(child2)
        noParent = []
        noParent2 = []

        parents2 = [s.replace(" ", "") for s in parents2]  # gets rid of space
        while parentTags:
            row = table.add_row().cells  # Adding a row and then adding data in it.
            row[0].text = parentTags[0]  # Adds the parentTag to the table
            noParent.append(parentTags[0])
            parentTags.remove(parentTags[0])  # Removes that tag after use

            if e < len(fullText2):  # as long as variable e is not higher than the lines in fullText2
                if fullText2[e] in filtered_L:  # filtered_L contains the parent tags without a child tag

                    noParent2.append(" ")
                    row[1].text = " "  # No child tag, so adds emopty string to that cell
                    e += 1

                elif fullText2[e] not in filtered_L:
                    if child2:
                        row[1].text = child2[0]  # Adds childTag to table
                        e += 1
                        noParent2.append(child2[0])
                        child2.remove(child2[0])  # Removed that tag from the list
        """
        while parentTags: # In case there are any more parent tags left in the list
            row = table.add_row().cells # Adding a row and then adding data in it.
            row[0].text = parentTags[0]
            parentTags.remove(parentTags[0])
        while child2: #This is for orphan tags, but not finished
            row = table.add_row().cells # Adding a row and then adding data in it.
            row[1].text = child2[0]
            child2.remove(child2[0])
        """
        # Make sure everything is cleared before the program gets the next document
        child2.clear()
        parentTags.clear()
        child.clear()
        report3.save('report3.docx')  # Saves in document "report3"

        # creates a dict for parent and child tags
        dicts = {}
        dicts = dict(zip(parents2, childCopy))  # creates a dictrionary if there is a child tag and parent tag
        # print(noParent)
        # print(noParent2)
        noParent = [s.replace(" ", "") for s in noParent]
        global dicts10
        if noParent2:
            dicts10 = dict(zip(noParent, noParent2))
        # print(dicts10)
        # print(parents2)
        global dicts3
        dicts3 = {}  # will hold parentTag and text, Orphan tags
        global dicts2
        dicts2 = {}  # will hold parentTag and text

        for x in parents2:  # creates dicttionary for child tags and text
            text2 = self.removeParent(everything)  # child tag and text
            # text8 = [s.replace(" ", "") for s in text2]
            text3 = self.removechild(text2)  # only text list
            text4 = self.removeText(text2)  # child tags
            # text8 = [s.replace(" ", "") for s in text4]
            dicts3 = dict(zip(parents2, text3))  # creates a dictionary with child tags and text
            sorted(dicts3.keys())  # sorts the keys in the dictionary

            # for x, y in dicts.items():
            # row = table.add_row().cells  # Adding a row and then adding data in it.
            #        row[0].text = x
            #       row[1].text = y
            # text1 = int(str(list(fullText)))
            # print(everything)
            text2 = self.removeParent(everything)  # child tag and text
            # text2 = removechild(everything)  # parent tags and text

            # print(text2)
            # text3 = removeParent(text2)  # only text list
            # text9 = ('"""' + str(text2) + '"""')  # child tag and text
            text3 = self.removechild(text2)  # only text list
            # print(text3)
            text4 = self.removeText(text2)  # child tags
            # print(text4) #only parent tag list
            # text7 = [s.replace(" ", "") for s in text3]
            text8 = [s.replace(" ", "") for s in text4]

            dicts2 = dict(zip(parents2, text3))  # creates a dictionary with child tags and text
            sorted(dicts2.keys())  # sorts the keys in the dictionary
            dicts2Copy.update(dicts2)

        m = 0
        k = 0
        while m <= len(parents2):
            if dicts2:
                for key, value in dicts2.items():

                    # for key, value in dicts2Copy.items() and key, value in dicts3.items(): #work on this here and try
                    report3.add_paragraph("\n")

                    report3.add_paragraph(key)
                    stringKey = str(key)
                    report3.add_paragraph(value)
                    if fullText2[k] not in filtered_L:  # check if it is an orphan tag
                        k += 1
                        if str(stringKey) in dicts10:
                            report3.add_paragraph(dicts10[str(stringKey)], style='List Bullet')
                            keyCheck = (dicts10[str(stringKey)].replace('[', ''))
                            keyCheck2 = (keyCheck.replace(']', ''))
                            keyCheck3 = (keyCheck2.replace(']', ''))
                            keyCheck4 = (keyCheck3.replace(' ', ''))
                            # print(keyCheck4)
                            if keyCheck4 in dicts2Copy:
                                report3.add_paragraph("       " + dicts2Copy[str(keyCheck4)])
                            m += 2
                        else:
                            m += 2
                            pass
                    else:
                        report3.add_paragraph(key + " is an orphan tag")
                        m += 2
                        k += 1

        report3.save('report3.docx')
        return dicts2Copy

    def removeParent(text):  # removes parent tags
        childAfter = []
        for line in text:
            childAfter = [i.rsplit('[', 1)[0] for i in text]  # removes parent tags
            childAfter = [re.sub("[\(\[].*?[\)\]]", "", e) for e in childAfter]  # removes parent tags that are left
            childAfter = [re.sub("[\{\[].*?[\)\}]", "", e) for e in childAfter]  # removes "pass", "fail", etc.
        return childAfter

    def removeText(text6):  # this should remove everything before the parent tag
        childAfter = [s.split(None, 1)[0] for s in text6]
        return childAfter

    def removeAfter(childtags):  # removes everything after the  tag, example "pass"
        seperator = ']'

        # for line in childtags:
        childAfter = [i.rsplit(']', 1)[0] + seperator for i in childtags]
        return childAfter

    def removechild(text):  # removes child, this one needs fixing
        mylst = []
        mylst = [s.split(None, 1)[1] for s in text]
        return mylst

    def getDirectory(self):
        response = QFileDialog.getExistingDirectory(self,
                                                    caption='Select a Folder')
        return response

    def getSaveFileName(self):
        fileFilter = 'Word Document (*.docx);; All File Types (*.*)'
        response = QFileDialog.getSaveFileName(parent=self,
                                               caption='Select a File',
                                               directory='report3.docx',
                                               filter=fileFilter,
                                               initialFilter='Word Document (*.docx)')
        print(response)
        return response[0]

    # This function will open up the report automatically
    def getDocument(self):
        if platform.system() == 'Darwin':
            subprocess.check_call(['open', 'report3.docx'])
        elif platform.system() == 'Windows':
            os.startfile('report3.docx')
        # os.startfile(report3) # try either one for windows if the first option gives error
        else:
            subprocess.call('xdg-open', report3)


if __name__ == '__main__':
    # Creates a word document, saves it as "report 3, and also adds a heading
    report3 = Document()
    report3.add_heading('Report', 0)  # create word document
    paragraph = report3.add_paragraph()
    report3.save('report3.docx')
    dicts2Copy = {}
    # This will hold the dicts2 content in all documents

    # declaring different lists that will be used to store, tags and sentences
    parentTags = []
    parent = []  # This will be used to store everything
    child = []  # Used to Store child tags
    noChild = []  # Used to Store parentTags with no child
    withChild = []  # Used to Store parentTags with child tag
    parents = []  # Will be used for future function

    app = QApplication(sys.argv)
    app.setStyleSheet('''QWidget{font-size: 35px;}''')
    myApp = MyApp()
    myApp.show()

    try:
        sys.exit(app.exec_())
    except SystemExit:
        print('Closing Program...')
